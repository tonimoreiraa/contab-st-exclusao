from docx import Document
from typing import List, Union, Dict
import pandas as pd

def edit_docx_table(
    input_path: str,
    output_path: str,
    table_index: int = 0,
    changes: Dict[tuple, str] = None
) -> None:
    doc = Document(input_path)
    
    # Get the specified table
    try:
        table = doc.tables[table_index]
    except IndexError:
        raise ValueError(f"No table found at index {table_index}")
    
    # Apply cell changes
    row_minus = 0
    if changes:
        for (row_idx, col_idx), new_text in changes.items():
            try:
                cell = table.cell(row_idx - row_minus, col_idx)
                cell.text = str(new_text)
            except IndexError:
                table_index = table_index + 1
                table = doc.tables[table_index]
                row_minus = row_minus + row_idx
                cell = table.cell(row_idx - row_minus, col_idx)
                cell.text = str(new_text)

                print(f"Warning: Cell ({row_idx}, {col_idx}) not found")
    
    # Save the modified document
    doc.save(output_path)

def get_csv_data(input_path, delimiter = ';'):
    df = pd.read_csv(input_path, delimiter=delimiter)
    data_rows = df.values.tolist()
    return data_rows

# Example usage
if __name__ == "__main__":
    changes = {
    }
    csv_data = get_csv_data('input.csv')
    for row in range(0, len(csv_data)):
        for col in range(0, len(csv_data[row])):
            changes[(row + 3, col)] = csv_data[row][col]

    edit_docx_table(
        input_path="input.docx",
        output_path="output.docx",
        table_index=1,
        changes=changes,
    )